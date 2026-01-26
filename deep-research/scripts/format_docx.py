#!/usr/bin/env python3
"""
Word 文档格式化脚本

Author: 翟星人

功能：
1. 标题居中，黑色字体，加粗
2. 在日期后插入目录
3. 一级标题前分页
4. 表格实线边框，不跨页断开
5. 日期居中
6. 图片说明小字居中
7. 1/2/3级标题加粗
8. 附录参考文献左对齐
"""

import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_page_break_before(paragraph):
    """在段落前添加分页符"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pageBreakBefore = OxmlElement('w:pageBreakBefore')
    pPr.insert(0, pageBreakBefore)


def set_table_border(table):
    """设置表格实线边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    
    tblBorders = OxmlElement('w:tblBorders')
    
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def keep_table_together(table):
    """保持表格不跨页断开"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                pPr = paragraph._p.get_or_add_pPr()
                keepNext = OxmlElement('w:keepNext')
                keepLines = OxmlElement('w:keepLines')
                pPr.append(keepNext)
                pPr.append(keepLines)


def keep_paragraph_together(paragraph):
    """保持段落不断开"""
    pPr = paragraph._p.get_or_add_pPr()
    keepNext = OxmlElement('w:keepNext')
    keepLines = OxmlElement('w:keepLines')
    pPr.append(keepNext)
    pPr.append(keepLines)


def set_heading_style(paragraph, level=1):
    """设置标题样式：黑色加粗"""
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(12)


def set_caption_style(paragraph):
    """设置图片说明样式：小字居中"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(80, 80, 80)


def is_image_caption(text, prev_has_image):
    """判断是否为图片说明"""
    if prev_has_image and text and len(text) < 100:
        # 必须以特定词开头才算图片说明
        if text.startswith("上图") or text.startswith("图：") or text.startswith("图:"):
            return True
    return False


def paragraph_has_image(paragraph):
    """检查段落是否包含图片"""
    for run in paragraph.runs:
        if run._element.xpath('.//w:drawing') or run._element.xpath('.//w:pict'):
            return True
    return False


def is_horizontal_rule(paragraph):
    """检查是否为分割线（文本或绘图元素）"""
    text = paragraph.text.strip()
    # 检查文本形式的分割线
    if text == "---" or text == "***" or text == "___" or (len(text) > 0 and all(c == '-' for c in text)):
        return True
    # 检查 pandoc 生成的绘图形式水平线（包含 line 或 rect 且文本为空，但不包含图片）
    if text == "":
        xml_str = paragraph._p.xml
        has_drawing = 'w:pict' in xml_str or 'w:drawing' in xml_str
        has_line = 'v:line' in xml_str or 'v:rect' in xml_str or '<a:ln' in xml_str
        has_image = 'a:blip' in xml_str or 'v:imagedata' in xml_str or 'r:embed' in xml_str
        # 只有有绘图、有线条、但没有图片时才是水平线
        if has_drawing and has_line and not has_image:
            return True
    return False


def is_reference_item(text):
    """判断是否为参考文献条目"""
    if re.match(r'^\d+\.\s', text):
        if 'http' in text or '地址' in text or 'github' in text.lower():
            return True
    return False


def add_toc_field(paragraph):
    """向段落添加目录域"""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.add_text('请右键点击此处，选择"更新域"以生成目录')
    run._r.append(fldChar3)


def format_docx(input_path, output_path=None):
    """格式化 Word 文档"""
    if output_path is None:
        output_path = input_path
    
    doc = Document(input_path)
    
    is_first_heading1 = True
    prev_was_code = False
    prev_has_image = False
    in_appendix = False
    date_para_index = -1
    paragraphs_to_remove = []
    
    # 第一遍遍历：找到并删除 Markdown 转换来的手写目录章节
    # 查找 Heading 2 样式且文本为"目录"的段落，以及紧随其后的 Compact 列表
    in_md_toc_section = False
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name is None:
            style_name = ""
        
        # 检测 Markdown 转换来的目录章节开始（Heading 2 且文本为"目录"）
        if "Heading 2" in style_name and text == "目录":
            in_md_toc_section = True
            paragraphs_to_remove.append(paragraph)
            continue
        
        # 在目录章节内，删除 Compact 样式的列表项（手写目录内容）
        if in_md_toc_section:
            # 遇到 Heading 样式，说明目录章节结束
            if "Heading" in style_name:
                in_md_toc_section = False
                continue
            # 删除 Compact 样式的列表项
            if style_name == "Compact":
                paragraphs_to_remove.append(paragraph)
                continue
            elif text == "":
                paragraphs_to_remove.append(paragraph)
                continue
            else:
                in_md_toc_section = False
    
    # 第二遍遍历：处理样式、标记删除分割线
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name is None:
            style_name = ""
        
        # 跳过已标记删除的段落
        if paragraph in paragraphs_to_remove:
            continue
        
        # 删除分割线段落
        if is_horizontal_rule(paragraph):
            paragraphs_to_remove.append(paragraph)
            continue
        
        # 检查是否进入附录部分
        if text.startswith("四、") or "附录" in text:
            in_appendix = True
        
        # 附录中的参考文献左对齐
        if in_appendix and is_reference_item(text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue
        
        # 检查当前段落是否有图片
        current_has_image = paragraph_has_image(paragraph)
        
        # 处理图片说明
        if is_image_caption(text, prev_has_image):
            set_caption_style(paragraph)
            prev_has_image = False
            continue
        
        prev_has_image = current_has_image
        
        # First Paragraph 样式应该左对齐（pandoc 可能设置为居中）
        if style_name == "First Paragraph":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # 处理文档标题
        if "Heading 1" in style_name or style_name == "Title":
            set_heading_style(paragraph, 1)
            
            if is_first_heading1 and ("调研报告" in text or i < 3):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                is_first_heading1 = False
                continue
        
        # 处理日期行，记录索引
        if "调研日期" in text or re.match(r'.*\d{4}年\d{1,2}月\d{1,2}日.*', text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_para_index = i
        
        # 一级标题前分页
        if "Heading" in style_name:
            if text.startswith("一、") or text.startswith("二、") or text.startswith("三、") or text.startswith("四、"):
                add_page_break_before(paragraph)
                set_heading_style(paragraph, 1)
            elif "Heading 2" in style_name:
                set_heading_style(paragraph, 2)
            elif "Heading 3" in style_name:
                set_heading_style(paragraph, 3)
            else:
                set_heading_style(paragraph, 1)
        
        # 代码块保持不断开
        if paragraph.style and "Code" in str(paragraph.style.name):
            keep_paragraph_together(paragraph)
            prev_was_code = True
        elif prev_was_code and text.startswith("```"):
            prev_was_code = False
    
    # 删除待删除的段落
    for paragraph in paragraphs_to_remove:
        p = paragraph._element
        parent = p.getparent()
        if parent is not None:
            parent.remove(p)
    
    # 插入 Word 目录域：在日期行后面插入
    target_para = None
    for i, p in enumerate(doc.paragraphs):
        if "调研日期" in p.text or re.match(r'.*\d{4}年\d{1,2}月\d{1,2}日.*', p.text):
            if i + 1 < len(doc.paragraphs):
                target_para = doc.paragraphs[i+1]
            break
    
    if target_para:
        toc_title = target_para.insert_paragraph_before("目录")
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in toc_title.runs:
            run.font.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        toc_field = target_para.insert_paragraph_before()
        add_toc_field(toc_field)
    
    # 最后一遍：修复 First Paragraph 样式的对齐（pandoc 默认居中）
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name == "First Paragraph":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 处理表格
    for table in doc.tables:
        set_table_border(table)
        keep_table_together(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    doc.save(output_path)
    print(f"格式化完成: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python format_docx.py <input.docx> [output.docx]")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    format_docx(input_file, output_file)
